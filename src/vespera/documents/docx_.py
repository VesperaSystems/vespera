"""DOCX text extraction via python-docx."""

from pathlib import Path

import docx


def extract_text(path: Path) -> str:
    document = docx.Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)
